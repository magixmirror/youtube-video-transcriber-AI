from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

def add_formatting(run, format_text):
    if '**' in format_text:
        parts = format_text.split('**')
        for i, part in enumerate(parts):
            if i % 2 == 1:  # Odd parts are within '**'
                run.add_text(part).bold = True
            else:
                run.add_text(part)
        format_text = ''
    # ... (other formatting code remains unchanged)
    return format_text

def generate_word_file(markdown_content, file_path):
    document = Document()

    # Add Markdown content to Word document
    for line in markdown_content.split('\n'):
        if line.startswith('# '):  # Heading 1
            heading_text = line[2:].strip()  # Remove leading/trailing spaces
            heading = document.add_heading(heading_text, level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line.startswith('## '):  # Heading 2
            heading_text = line[3:].strip()  # Remove leading/trailing spaces
            heading = document.add_heading(heading_text, level=2)
        else:
            paragraph = document.add_paragraph()
            run = paragraph.add_run()
            formatted_text = add_formatting(run, line)
            if line.startswith('#'):
                # Add appropriate Markdown heading syntax
                heading_level = line.count('#')
                heading.text = formatted_text
                heading.style.font.size = Pt(14 - heading_level * 2)  # Adjust font size for different heading levels
            else:
                run.text = formatted_text
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Set default alignment for non-heading paragraphs

    # Save Word document
    document.save(file_path)

